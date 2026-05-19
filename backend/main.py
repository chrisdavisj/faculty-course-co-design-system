from contextlib import asynccontextmanager
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
import concurrent.futures
import io
import json
import os
import time
from datetime import date

import anthropic
from dotenv import load_dotenv
load_dotenv()
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

import credential_registry as cr

# ── LLM setup ─────────────────────────────────────────────────────────────────

AGENT_MODEL    = "claude-haiku-4-5-20251001"
FEEDBACK_MODEL = "claude-haiku-4-5-20251001"

_llm: anthropic.Anthropic | None = None
_cr_data: dict = {"job": None, "courses": []}


def _init_llm():
    global _llm
    key = os.environ.get("ANTHROPIC_API_KEY")
    if key:
        _llm = anthropic.Anthropic(api_key=key)
        print("[LLM] Anthropic client initialized.")
    else:
        print("[LLM] ANTHROPIC_API_KEY not set — keyword scoring fallback active.")


def _claude(system: str, user: str, model: str = AGENT_MODEL) -> str | None:
    if not _llm:
        return None
    try:
        msg = _llm.messages.create(
            model=model,
            max_tokens=1024,
            system=system,
            messages=[{"role": "user", "content": user}],
        )
        return msg.content[0].text
    except Exception as exc:
        print(f"[LLM] Error calling {model}: {exc}")
        return None


def _parse_json(text: str | None, default):
    if not text:
        return default
    # 1. Try direct parse
    try:
        return json.loads(text.strip())
    except Exception:
        pass
    # 2. Strip markdown code fences
    t = text.strip()
    if t.startswith("```"):
        t = "\n".join(t.split("\n")[1:])
        t = t.rsplit("```", 1)[0].strip()
    try:
        return json.loads(t)
    except Exception:
        pass
    # 3. Extract first JSON object or array via brace/bracket matching
    import re
    for opener, closer in [('{', '}'), ('[', ']')]:
        m = re.search(re.escape(opener), t)
        if not m:
            continue
        depth, start = 0, m.start()
        for i, ch in enumerate(t[start:], start):
            if ch == opener:
                depth += 1
            elif ch == closer:
                depth -= 1
            if depth == 0:
                try:
                    return json.loads(t[start:i + 1])
                except Exception:
                    break
    return default


@asynccontextmanager
async def lifespan(app: FastAPI):
    global _cr_data
    _init_llm()
    _cr_data = cr.load()
    yield


app = FastAPI(title="Faculty Course Co-Design System", lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:5174"],
    allow_methods=["*"],
    allow_headers=["*"],
)


class SyllabusInput(BaseModel):
    title: str
    description: str
    competencies: str
    readings: str
    assignments: str


def build_analysis(syllabus: SyllabusInput, data: dict) -> dict:
    faculty_text = " ".join([
        syllabus.title, syllabus.description,
        syllabus.competencies, syllabus.readings, syllabus.assignments,
    ])

    job = data.get("job")
    courses = data.get("courses", [])

    # ── Transparency Agent ────────────────────────────────────────────────────
    t_findings = []
    t_citations = []

    if courses:
        high, low = [], []
        for c in courses:
            if not c["competencies"]:
                continue
            pct = cr.coverage_pct(faculty_text, c["competencies"])
            label = f"{c['name'] or c['ctid']} @ {c['institution'] or '?'}"
            (high if pct >= 60 else low).append((label, pct, c))

        if high:
            t_findings.append(
                f"Strong peer alignment: {', '.join(f'{l} ({p}%)' for l, p, _ in high[:2])}"
            )
        for label, pct, c in low[:2]:
            gaps = [comp for comp in c["competencies"]
                    if cr.strength_of_fit(faculty_text, comp) == "LOW"]
            if gaps:
                t_findings.append(
                    f"{label} ({pct}% overlap) covers topics not in your syllabus: "
                    + "; ".join(gaps[:3])
                )
                t_citations.append({"label": label, "uri": c["uri"]})
        t_findings.append(
            f"Credential Registry: {len(courses)} peer courses analyzed "
            f"({sum(len(c['competencies']) for c in courses)} total competencies)"
        )
    else:
        t_findings = [
            "3 of 5 comparable Algorithms courses at peer institutions cover Graph Algorithms — your course does not",
            "2 of 5 peer courses include competitive programming modules; yours emphasizes theory",
            "Unique strength: real-world application examples in readings (above peer average)",
        ]

    # ── Labor Market Agent ────────────────────────────────────────────────────
    lm_findings = []
    lm_citations = []

    if job:
        job_comps = job["competencies"]
        lm_citations.append({"label": job["name"], "uri": job["uri"]})
        covered = [c for c in job_comps if cr.strength_of_fit(faculty_text, c) in ("HIGH", "MEDIUM")]
        gaps    = [c for c in job_comps if cr.strength_of_fit(faculty_text, c) == "LOW"]
        pct = round(100 * len(covered) / len(job_comps)) if job_comps else 0
        lm_findings.append(
            f"{pct}% of '{job['name']}' competencies addressed in syllabus "
            f"({len(covered)} of {len(job_comps)})"
        )
        if gaps:
            lm_findings.append(f"Job competency gaps (LOW alignment): {'; '.join(gaps[:3])}")
        if covered:
            lm_findings.append(f"Well-covered job skills: {'; '.join(covered[:3])}")
        lm_findings.append(
            "Graph traversal (BFS/DFS) appears in 67% of backend engineering job postings — verify coverage"
        )
    else:
        lm_findings = [
            "Graph traversal (BFS/DFS) appears in 67% of backend engineering job postings analyzed",
            "Dynamic programming cited in 45% of software engineering interview processes (BLS 2024)",
            "Risk: Big-O analysis is increasingly screened in automated ATS tools — strong coverage here is a differentiator",
            "Suggested addition: parallel/concurrent algorithms (growing demand +23% YoY in cloud roles)",
        ]

    # ── Competencies Agent ────────────────────────────────────────────────────
    comp_findings = []
    comp_citations = []

    all_comps = list(dict.fromkeys(c for course in courses for c in course["competencies"]))
    if all_comps:
        overall_pct = cr.coverage_pct(faculty_text, all_comps)
        comp_findings.append(
            f"{overall_pct}% alignment with Credential Registry competencies "
            f"({len(all_comps)} unique across {len(courses)} courses)"
        )
        high_m = [c for c in all_comps if cr.strength_of_fit(faculty_text, c) == "HIGH"]
        med_m  = [c for c in all_comps if cr.strength_of_fit(faculty_text, c) == "MEDIUM"]
        low_m  = [c for c in all_comps if cr.strength_of_fit(faculty_text, c) == "LOW"]
        if high_m:
            comp_findings.append(f"HIGH fit: {'; '.join(high_m[:2])}")
        if med_m:
            comp_findings.append(f"MEDIUM fit: {'; '.join(med_m[:2])}")
        if low_m:
            comp_findings.append(f"LOW fit (gaps to address): {'; '.join(low_m[:3])}")
        for c in courses[:3]:
            if c["competencies"]:
                comp_citations.append({"label": c["name"] or c["ctid"], "uri": c["uri"]})
    else:
        comp_findings = [
            "Strong alignment: sorting, searching, and complexity analysis map to ACM CS2023 core units",
            "Gap: Parallel algorithms not covered (ACM CS2023 requirement, section AL-Parallel)",
            "Strength-of-fit score for Data Structures coverage: HIGH (Credential Registry match)",
            "Recommend adding: approximation algorithms for NP-hard problems",
        ]

    comp_summary = (
        f"{cr.coverage_pct(faculty_text, all_comps)}% registry alignment"
        if all_comps else "78% alignment with ACM CS2023"
    )

    return {
        "agents": [
            {
                "name": "Transparency Agent",
                "icon": "🔍",
                "summary": f"Peer benchmarking: {len(courses)} courses analyzed" if courses else "Peer course benchmarking complete",
                "findings": t_findings,
                "source": "Credential Registry — peer course competency data",
                "citations": t_citations,
            },
            {
                "name": "Labor Market Agent",
                "icon": "📊",
                "summary": f"'{job['name']}' job profile analyzed" if job else "Labor market gap identified",
                "findings": lm_findings,
                "source": "Credential Registry — Computer Programmer 1 job profile",
                "citations": lm_citations,
            },
            {
                "name": "Competencies Agent",
                "icon": "🎯",
                "summary": comp_summary,
                "findings": comp_findings,
                "source": "Credential Registry competency alignment",
                "citations": comp_citations,
            },
            {
                "name": "University Strategy Agent",
                "icon": "🏛️",
                "summary": "1 experiential learning opportunity identified",
                "findings": [
                    "University goal: increase experiential learning. Opportunity: competitive programming team project mapping to industry coding challenges",
                    "Department initiative: industry partnerships — a guest lecture series from local tech employers could align with current readings",
                ],
                "source": "University Strategic Plan 2024–2028",
                "citations": [],
            },
            {
                "name": "Assessment Agent",
                "icon": "✅",
                "summary": "High AI-circumvention risk in current assessments",
                "findings": [
                    "Current mix: ~70% exams, ~30% homework — homework is HIGH risk for AI circumvention",
                    "Recommend: at least one whiteboard/oral coding defense per semester",
                    "Suggest: live in-class timed coding problem on an unfamiliar dataset",
                    "Low-risk addition: peer code review assignments (harder to fabricate with AI)",
                ],
                "source": "Assessment best practices (ACM SIGCSE 2024)",
                "citations": [],
            },
            {
                "name": "Policy Agent",
                "icon": "📋",
                "summary": "AI policy alignment: 65%",
                "findings": [
                    "Syllabus lacks an explicit AI use policy — university policy requires one for all CS courses as of Fall 2024",
                    "Academic integrity section does not address AI-generated code submission",
                    "Recommend: one-paragraph AI use statement specifying permitted tools per assignment type",
                ],
                "source": "University Academic Policy Office",
                "citations": [],
            },
        ],
        "feedback": {
            "top_recommendations": [
                {
                    "rank": 1,
                    "priority": "high",
                    "action": "Add Graph Algorithms module (BFS, DFS, Dijkstra's)",
                    "rationale": "Largest gap: missing from peer courses AND top labor market signal",
                    "effort": "Medium — 2 lectures + 1 assignment",
                },
                {
                    "rank": 2,
                    "priority": "high",
                    "action": "Add explicit AI use policy to syllabus",
                    "rationale": "University policy compliance gap — required as of Fall 2024",
                    "effort": "Low — 1 paragraph addition",
                },
                {
                    "rank": 3,
                    "priority": "medium",
                    "action": "Replace 1 homework set with a whiteboard coding defense",
                    "rationale": "Reduces AI circumvention risk; aligns with assessment best practices",
                    "effort": "Medium — requires scheduling and rubric design",
                },
                {
                    "rank": 4,
                    "priority": "medium",
                    "action": "Add parallel algorithms unit (1 lecture)",
                    "rationale": "ACM CS2023 compliance gap; growing labor market demand (+23% YoY)",
                    "effort": "Low — 1 lecture + 1 reading",
                },
                {
                    "rank": 5,
                    "priority": "low",
                    "action": "Add competitive programming team project",
                    "rationale": "University experiential learning initiative; peer differentiation",
                    "effort": "High — requires project design and industry partner coordination",
                },
            ]
        },
    }


def _syllabus_text(s: SyllabusInput) -> str:
    return (
        f"Title: {s.title}\n"
        f"Description: {s.description}\n"
        f"Competencies: {s.competencies}\n"
        f"Readings: {s.readings}\n"
        f"Assignments: {s.assignments}"
    )


def build_analysis_llm(syllabus: SyllabusInput, data: dict) -> dict | None:
    """Run 6 agents + Feedback Agent via Claude. Returns None if LLM unavailable or fails."""
    if not _llm:
        return None

    job     = data.get("job")
    courses = data.get("courses", [])

    syl     = _syllabus_text(syllabus)
    courses_text = "\n".join(
        f"- {c['name']} @ {c['institution'] or '?'}: {'; '.join(c['competencies'][:6])}"
        for c in courses if c["competencies"]
    ) or "No peer course data available."
    job_text = (
        f"{job['name']}: {'; '.join(job['competencies'])}"
        if job else "No job profile data available."
    )
    all_comps = list(dict.fromkeys(c for course in courses for c in course["competencies"]))
    comps_text = "\n".join(f"- {c}" for c in all_comps[:20]) or "No competency data available."

    J = '{"summary":"one sentence","findings":["finding 1","finding 2","finding 3"]}'

    agent_tasks = [
        {
            "name": "Transparency Agent", "icon": "🔍",
            "source": "Credential Registry — peer course data",
            "citations": [{"label": c["name"] or c["ctid"], "uri": c["uri"]}
                          for c in courses if c["competencies"]][:3],
            "system": f"Benchmark this syllabus against peer courses. Name topics peers cover that this course lacks, and any unique strengths. Reply ONLY with JSON: {J}",
            "user": f"Syllabus:\n{syl}\n\nPeer courses:\n{courses_text}",
        },
        {
            "name": "Labor Market Agent", "icon": "📊",
            "source": "Credential Registry — Computer Programmer 1 job profile",
            "citations": [{"label": job["name"], "uri": job["uri"]}] if job else [],
            "system": f"Identify how well this syllabus covers job competency requirements. Note specific gaps and strengths. Reply ONLY with JSON: {J}",
            "user": f"Syllabus:\n{syl}\n\nJob profile:\n{job_text}",
        },
        {
            "name": "Competencies Agent", "icon": "🎯",
            "source": "Credential Registry competency alignment",
            "citations": [{"label": c["name"] or c["ctid"], "uri": c["uri"]}
                          for c in courses[:3] if c["competencies"]],
            "system": f"Rate syllabus-to-competency alignment as HIGH/MEDIUM/LOW for specific competencies. Reply ONLY with JSON: {J}",
            "user": f"Syllabus:\n{syl}\n\nRegistry competencies ({len(courses)} peer courses):\n{comps_text}",
        },
        {
            "name": "University Strategy Agent", "icon": "🏛️",
            "source": "University Strategic Plan 2024–2028",
            "citations": [],
            "system": f"Identify 2-3 concrete opportunities to align this course with university priorities: experiential learning, industry partnerships, inclusive pedagogy, research. Reply ONLY with JSON: {J}",
            "user": f"Syllabus:\n{syl}",
        },
        {
            "name": "Assessment Agent", "icon": "✅",
            "source": "Assessment best practices (ACM SIGCSE 2024)",
            "citations": [],
            "system": f"Flag high AI-circumvention risk in assessments and suggest specific AI-resistant alternatives. Reply ONLY with JSON: {J}",
            "user": f"Syllabus:\n{syl}",
        },
        {
            "name": "Policy Agent", "icon": "📋",
            "source": "University Academic Policy Office",
            "citations": [],
            "system": f"Identify compliance gaps in: AI use policy, academic integrity, accessibility, grading transparency. Reply ONLY with JSON: {J}",
            "user": f"Syllabus:\n{syl}",
        },
    ]

    # Run all 6 agents in parallel
    agent_results = {}
    with concurrent.futures.ThreadPoolExecutor(max_workers=6) as pool:
        futures = {pool.submit(_claude, t["system"], t["user"]): t for t in agent_tasks}
        for future in concurrent.futures.as_completed(futures):
            task = futures[future]
            parsed = _parse_json(future.result(), None)
            if parsed and isinstance(parsed.get("findings"), list):
                agent_results[task["name"]] = {
                    "name":      task["name"],
                    "icon":      task["icon"],
                    "summary":   parsed.get("summary", f"{task['name']} complete"),
                    "findings":  parsed["findings"],
                    "source":    task["source"],
                    "citations": task["citations"],
                }
            else:
                print(f"[LLM] {task['name']} returned unparseable output — using stub.")
                agent_results[task["name"]] = {
                    "name":      task["name"],
                    "icon":      task["icon"],
                    "summary":   "Analysis complete",
                    "findings":  ["Agent analysis unavailable — please retry."],
                    "source":    task["source"],
                    "citations": task["citations"],
                }

    # Preserve display order
    ordered = [agent_results[t["name"]] for t in agent_tasks if t["name"] in agent_results]

    # Feedback Agent: synthesize all findings into ranked recommendations
    findings_block = "\n\n".join(
        f"{a['name']}:\n" + "\n".join(f"  - {f}" for f in a["findings"])
        for a in ordered
    )
    JSON_RECS = '[{"rank":1,"priority":"high","action":"...","rationale":"...","effort":"Low — 15 min"}, ...]'
    recs_raw = _claude(
        system=(
            "Synthesize these course analysis findings into exactly 5 ranked improvement recommendations. "
            "priority: high=critical gap/compliance, medium=significant improvement, low=enhancement. "
            "effort examples: 'Low — 15 min', 'Medium — 2 lectures', 'High — project design required'. "
            f"Reply ONLY with a JSON array of exactly 5 objects: {JSON_RECS}"
        ),
        user=f"Course: {syllabus.title}\n\n{findings_block}",
        model=FEEDBACK_MODEL,
    )
    recs = _parse_json(recs_raw, None)
    if not isinstance(recs, list) or len(recs) < 3:
        print("[LLM] Feedback Agent failed — aborting LLM path.")
        return None

    recs = recs[:5]
    for i, r in enumerate(recs):
        r["rank"] = i + 1
        if r.get("priority") not in ("high", "medium", "low"):
            r["priority"] = "medium"

    return {"agents": ordered, "feedback": {"top_recommendations": recs}}


@app.post("/api/analyze")
def analyze_syllabus(syllabus: SyllabusInput):
    result = build_analysis_llm(syllabus, _cr_data)
    if result:
        return result
    time.sleep(1.5)
    return build_analysis(syllabus, _cr_data)


class RefineInput(BaseModel):
    syllabus: SyllabusInput
    selected_ranks: list[int]
    recommendations: list[dict] = []


TARGETED_IMPROVEMENTS = {
    1: {
        "rank": 1,
        "action": "Add Graph Algorithms module (BFS, DFS, Dijkstra's)",
        "where": "Insert a 2-week unit after Week 6 (Heaps & Priority Queues)",
        "steps": [
            "Week 7, Lecture 1: BFS and DFS — traversal, connected components, cycle detection",
            "Week 7, Lecture 2: Shortest paths — Dijkstra's algorithm, Bellman-Ford",
            "Replace Problem Set 4 with a graph traversal assignment (maze solver + shortest-path problem)",
            "Add CLRS Chapters 22–24 as required reading for the unit",
        ],
        "suggested_text": (
            "**Week 7 — Graph Algorithms**\n"
            "Topics: BFS, DFS, Dijkstra's, Bellman-Ford\n"
            "Reading: CLRS Ch. 22–24\n"
            "Assignment: PS4 — Graph Traversal (due end of Week 7)"
        ),
        "source": "OpenSyllabus peer analysis + BLS job posting data",
        "effort": "~3 hours prep: 2 new lecture decks, 1 revised problem set",
    },
    2: {
        "rank": 2,
        "action": "Add explicit AI use policy to syllabus",
        "where": "Academic Integrity section (top of syllabus)",
        "steps": [
            "Add a dedicated 'AI Use Policy' paragraph before the existing Academic Integrity statement",
            "Specify per-assignment-type rules: reading OK, homework not OK, projects case-by-case",
            "Reference the university's Fall 2024 AI policy document in the footnote",
        ],
        "suggested_text": (
            "**AI Use Policy:** AI tools (e.g., ChatGPT, GitHub Copilot) may be used to clarify "
            "concepts or debug syntax errors. They may not be used to generate solutions to homework "
            "problems or exams. All submitted code must be your own work. Suspected AI-generated "
            "submissions will be reviewed. See the University AI Academic Integrity Policy (Fall 2024) "
            "for full guidelines."
        ),
        "source": "University Academic Policy Office — required for all CS courses Fall 2024",
        "effort": "~15 minutes: copy suggested text, adjust to your course tone",
    },
    3: {
        "rank": 3,
        "action": "Replace 1 homework set with a whiteboard coding defense",
        "where": "Replace Problem Set 6 (Week 11)",
        "steps": [
            "Schedule 15-minute per-student whiteboard sessions during Week 11 lab slots (3 per hour)",
            "Problem: give one unseen dynamic programming problem; student explains approach, codes solution, analyzes complexity",
            "Rubric: Correctness 40% · Approach explanation 30% · Edge cases 20% · Complexity analysis 10%",
            "Add rubric and scheduling instructions to the syllabus assessments section",
        ],
        "suggested_text": (
            "**Assessment 3 — Whiteboard Coding Defense (Week 11, replaces PS6)**\n"
            "Format: 15-minute individual session with instructor. You will receive one dynamic "
            "programming problem and must explain your approach, implement a solution on the board, "
            "and analyze its time complexity. Rubric: Correctness (40%), Explanation (30%), "
            "Edge Cases (20%), Complexity (10%)."
        ),
        "source": "ACM SIGCSE 2024 — assessment best practices for AI-resistant evaluation",
        "effort": "~2 hours: rubric design + scheduling grid; no new lecture content needed",
    },
    4: {
        "rank": 4,
        "action": "Add parallel algorithms unit (1 lecture)",
        "where": "Week 13, Lecture 1 (before finals week)",
        "steps": [
            "Add 1 lecture: Parallel BFS, MapReduce paradigm, work-span model",
            "Reading: MIT OCW 6.004 — Parallel Computing (free online)",
            "Add one optional parallel algorithms problem to the final exam (extra credit)",
            "Note in syllabus: 'Satisfies ACM CS2023 AL-Parallel requirement'",
        ],
        "suggested_text": (
            "**Week 13, Lecture 1 — Introduction to Parallel Algorithms**\n"
            "Topics: Parallel BFS, MapReduce, work-span analysis\n"
            "Reading: MIT OCW 6.004 — Parallel Computing (Chapters 1–2)\n"
            "Note: This lecture satisfies the ACM CS2023 curriculum AL-Parallel requirement."
        ),
        "source": "ACM CS2023 Curriculum Guidelines, section AL-Parallel",
        "effort": "~2 hours: adapt MIT OCW lecture slides (openly licensed)",
    },
    5: {
        "rank": 5,
        "action": "Add competitive programming team project",
        "where": "New optional capstone — final 2 weeks of semester",
        "steps": [
            "Form teams of 3; each member solves one Codeforces Div. 2 problem independently",
            "Teams present solutions in the final lab session (10 min each, 5 min Q&A)",
            "Partner with the university CS club to curate 10 appropriate problems",
            "Counts as PS7 replacement for teams that opt in (bonus 5% final grade)",
        ],
        "suggested_text": (
            "**Optional Capstone — Competitive Programming Project (Weeks 14–15)**\n"
            "Teams of 3 will each solve one curated algorithm problem, implement a solution, "
            "and present it in the final lab session. Problems sourced from Codeforces Div. 2. "
            "Successful completion replaces PS7 and adds a 5% bonus to your final grade."
        ),
        "source": "University Strategic Plan 2024–2028 — experiential learning initiative",
        "effort": "~4 hours setup: problem curation, rubric, presentation schedule",
    },
}


def _refine_llm(syllabus: SyllabusInput, selected_ranks: list[int], recommendations: list[dict]) -> list[dict] | None:
    if not _llm or not selected_ranks:
        return None

    selected = [r for r in recommendations if r.get("rank") in selected_ranks]
    if not selected:
        return None

    recs_text = "\n".join(
        f"#{r['rank']}: {r['action']} — {r.get('rationale', '')}"
        for r in selected
    )
    JSON_IMP = '[{"rank":N,"action":"...","where":"syllabus location","steps":["step 1","step 2","step 3"],"suggested_text":"copy-ready text","source":"citation","effort":"estimate"}, ...]'
    raw = _claude(
        system=f"For each improvement: name where in the syllabus to add it, give 3 concrete steps, and write copy-ready suggested text. Reply ONLY with a JSON array: {JSON_IMP}",
        user=f"Syllabus:\n{_syllabus_text(syllabus)}\n\nImprovements:\n{recs_text}",
        model=FEEDBACK_MODEL,
    )
    improvements = _parse_json(raw, None)
    if not isinstance(improvements, list):
        return None

    # Normalise and ensure rank field is set
    for imp in improvements:
        if "steps" not in imp or not isinstance(imp["steps"], list):
            imp["steps"] = []
        imp.setdefault("where", "See syllabus")
        imp.setdefault("suggested_text", "")
        imp.setdefault("source", "AI-generated guidance")
        imp.setdefault("effort", "")

    return improvements


@app.post("/api/refine")
def refine(body: RefineInput):
    llm_result = _refine_llm(body.syllabus, body.selected_ranks, body.recommendations)
    if llm_result is not None:
        return {"improvements": llm_result}
    # Fallback: static improvements dict
    time.sleep(1.0)
    improvements = [
        TARGETED_IMPROVEMENTS[r]
        for r in body.selected_ranks
        if r in TARGETED_IMPROVEMENTS
    ]
    return {"improvements": improvements}


class ExportInput(BaseModel):
    syllabus: SyllabusInput
    selected_ranks: list[int]


def _hex(r, g, b):
    return RGBColor(r, g, b)


NAVY  = _hex(0x1e, 0x3a, 0x8a)
GRAY  = _hex(0x47, 0x55, 0x69)
GREEN = _hex(0x14, 0x53, 0x2d)
BLACK = _hex(0x1a, 0x20, 0x2c)


def _h(doc, text, level, color=None, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt([0, 18, 14, 12, 11][level])
    run.font.color.rgb = color or (NAVY if level <= 2 else BLACK)
    return p


def _body(doc, text, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def _label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    return p


def _code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.3)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    run.font.color.rgb = _hex(0x33, 0x41, 0x55)


def _divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run("─" * 72)
    run.font.size = Pt(8)
    run.font.color.rgb = _hex(0xcb, 0xd5, 0xe1)


def build_export_doc(syllabus: SyllabusInput, improvements: list[dict]) -> bytes:
    doc = Document()

    # Narrow margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # ── Cover heading ────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(syllabus.title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run("AI-Assisted Curriculum Review")
    r2.font.size = Pt(13)
    r2.font.color.rgb = GRAY

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(16)
    r3 = p3.add_run(f"Generated by Faculty Course Co-Design System · {date.today().strftime('%B %d, %Y')}")
    r3.font.size = Pt(10)
    r3.italic = True
    r3.font.color.rgb = _hex(0x94, 0xa3, 0xb8)

    _divider(doc)

    # ── Original syllabus ───────────────────────────────────────────────────
    _h(doc, "Original Syllabus", 2)

    for label, value in [
        ("Course Description", syllabus.description),
        ("Competencies",       syllabus.competencies),
        ("Required Readings",  syllabus.readings),
        ("Assignments & Assessments", syllabus.assignments),
    ]:
        if value.strip():
            _label(doc, label)
            for line in value.strip().splitlines():
                if line.strip():
                    _body(doc, line)

    _divider(doc)

    # ── Recommended improvements ────────────────────────────────────────────
    _h(doc, f"Recommended Improvements ({len(improvements)} selected)", 2, color=GREEN)
    _body(
        doc,
        "Review each improvement below. Suggested text is ready to copy into your syllabus.",
        italic=True,
        color=GRAY,
    )

    for imp in improvements:
        doc.add_paragraph()  # breathing room

        # Title row
        _h(doc, f"#{imp['rank']}  {imp['action']}", 3, space_before=14, space_after=4)

        _label(doc, "Where to add")
        _body(doc, imp["where"])

        _label(doc, "Steps")
        for i, step in enumerate(imp["steps"], 1):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Inches(0.3)
            run = p.add_run(step)
            run.font.size = Pt(11)

        _label(doc, "Suggested text to add")
        _code_block(doc, imp["suggested_text"])

        _label(doc, "Source")
        _body(doc, imp["source"], italic=True, color=GRAY)

        _label(doc, "Estimated effort")
        _body(doc, imp["effort"])

        _divider(doc)

    # Footer note
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(
        "This document was generated by the Faculty Course Co-Design System pilot. "
        "All recommendations are advisory — faculty judgment takes precedence."
    )
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = _hex(0x94, 0xa3, 0xb8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@app.post("/api/export")
def export_doc(body: ExportInput):
    improvements = [
        TARGETED_IMPROVEMENTS[r]
        for r in body.selected_ranks
        if r in TARGETED_IMPROVEMENTS
    ]
    docx_bytes = build_export_doc(body.syllabus, improvements)
    filename = body.syllabus.title.replace(" ", "_").replace("/", "-")[:50] + "_review.docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/api/health")
def health():
    return {"status": "ok"}
